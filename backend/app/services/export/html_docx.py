"""叙事类导出：HTML / DOCX"""
import io


def to_html(article, body_text: str) -> str:
    body_html = body_text.replace("\n", "<br>")
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{article.title or article.topic}</title>
<style>
body {{ font-family: "PingFang SC", "Microsoft YaHei", sans-serif; margin: 2em; line-height: 1.6; }}
h1 {{ font-size: 1.5em; }}
h2 {{ font-size: 1.2em; margin-top: 1.5em; }}
.content {{ margin-top: 1em; }}
</style>
</head>
<body>
<h1>{article.title or article.topic}</h1>
<p><em>主题：{article.topic} | 形式：{article.content_format} | 平台：{article.platform}</em></p>
<hr>
<div class="content">{body_html}</div>
</body>
</html>"""


def to_docx(article, parts: list[tuple[str, str]]) -> tuple[bytes, str]:
    from docx import Document
    doc = Document()
    doc.add_heading(article.title or article.topic or "未命名", 0)
    doc.add_paragraph(f"主题：{article.topic} | 形式：{article.content_format} | 平台：{article.platform}")
    for title, body in parts:
        doc.add_heading(title, level=1)
        for para in body.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue(), f"{(article.topic or 'article').replace('/', '-')}.docx"
